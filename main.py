import gradio as gr
import pandas as pd
import time
import json
from datetime import datetime
import io
import cv2
import numpy as np
from ultralytics import YOLO
import os
import tempfile
import shutil
import uuid

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not installed. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF

    PDF_AVAILABLE = True
except ImportError:
    print("⚠️ fpdf2 not installed. Install with: pip install fpdf2")
    PDF_AVAILABLE = False

# Load your trained YOLOv8 model
MODEL_PATH = r'C:EO-IR video classification system\weights\best.pt'

# Try to load model, fallback to mock if not available
try:
    model = YOLO(MODEL_PATH)
    MODEL_LOADED = True
    print("✅ YOLOv8 model loaded successfully!")
    print(f"🤖 Model classes: {model.names}")
except Exception as e:
    print(f"⚠️ Could not load model: {e}")
    print("🔄 Using mock detections for demo purposes")
    MODEL_LOADED = False

# Military/Surveillance object categories (match your actual training classes)
DETECTION_CATEGORIES = [
    'Artillery', 'Missile', 'Radar', 'M. Rocket Launcher',
    'Soldier', 'Tank', 'Vehicle'
]

# Map your model's class names to display names (based on your actual model output)
CLASS_NAME_MAPPING = {
    0: 'Artillery',  
    1: 'Missile',
    2: 'Radar',
    3: 'M. Rocket Launcher',
    4: 'Soldier',
    5: 'Tank',
    6: 'Vehicle',
}


def assess_threat_level(object_type):
    """Assess threat level based on detected object type"""
    threat_mapping = {
        'Artillery': 'Critical',
        'Missile': 'Critical',
        'Radar': 'High',
        'M. Rocket Launcher': 'Critical',
        'Soldier': 'Medium',
        'Tank': 'High',
        'Vehicle': 'Medium'
    }
    return threat_mapping.get(object_type, 'Low')


def convert_frame_to_timestamp(frame_number, fps):
    """Convert frame number to timestamp format"""
    total_seconds = int(frame_number / fps)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def analyze_surveillance_footage(video_file, selected_objects, confidence_threshold):
    """Analyze uploaded surveillance footage for military objects"""
    if not video_file:
        return (
            " Please upload a video file first.",
            pd.DataFrame(),
            "",
            gr.update(visible=False),
            None,  # Original video
            None  # Annotated video
        )

    if not selected_objects:
        return (
            " Please select at least one object type to detect.",
            pd.DataFrame(),
            "",
            gr.update(visible=False),
            video_file,  # Original video
            None  # Annotated video
        )

    try:
        print(f"🔄 Starting analysis of {video_file.name}...")

        if not MODEL_LOADED:
            return (
                "YOLOv8 model not loaded. Check model path.",
                pd.DataFrame(),
                "",
                gr.update(visible=False),
                video_file,  
                None  
            )

        # Create temporary directory for output
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, f"annotated_{uuid.uuid4().hex[:8]}.mp4")

        cap = cv2.VideoCapture(video_file.name)
        if not cap.isOpened():
            return (
                " Could not open video file.",
                pd.DataFrame(),
                "",
                gr.update(visible=False),
                video_file,  # Original video
                None  # Annotated video
            )

        # Get video properties
        fps = cap.get(cv2.CAP_PROP_FPS)
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

        print(f"📹 Video: {width}x{height}, {fps:.1f} FPS, {total_frames} frames")

        # Create video writer with H.264 codec for better compatibility
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))

        detections = []
        frame_count = 0

        print("🎬 Processing frames with YOLOv8...")

        while True:
            ret, frame = cap.read()
            if not ret:
                break

            frame_count += 1

            # Run YOLO inference with specified confidence threshold
            results = model(frame, conf=confidence_threshold / 100, verbose=False)

            # Create annotated frame
            annotated_frame = frame.copy()

            if results and len(results) > 0 and results[0].boxes is not None:
                # Use YOLO's built-in plotting for professional annotations
                annotated_frame = results[0].plot(
                    line_width=2,
                    font_size=1,
                    labels=True,
                    boxes=True,
                    conf=True
                )

                # Record detections for the table
                for box in results[0].boxes:
                    class_id = int(box.cls[0])
                    confidence = float(box.conf[0])
                    class_name = model.names.get(class_id, f"class_{class_id}")

                    # Map to display name
                    display_name = CLASS_NAME_MAPPING.get(class_name, class_name)

                    # Filter by selected objects
                    if display_name in selected_objects:
                        current_time = convert_frame_to_timestamp(frame_count, fps)
                        x1, y1, x2, y2 = box.xyxy[0].cpu().numpy()

                        detections.append({
                            'object': display_name,
                            'timestamp': current_time,
                            'confidence': f"{confidence * 100:.1f}%",
                            'coordinates': f"({int(x1)}, {int(y1)}, {int(x2 - x1)}, {int(y2 - y1)})",
                            'count': 1,
                            'threat_level': assess_threat_level(display_name),
                            'duration': '00:00:01'
                        })

                        print(f"🎯 Detected: {display_name} at {confidence:.2f} confidence")

            # Write the annotated frame
            out.write(annotated_frame)

            # Progress indicator
            if frame_count % 30 == 0:
                progress = (frame_count / total_frames) * 100 if total_frames > 0 else 0
                print(f"📊 Processed {frame_count}/{total_frames} frames ({progress:.1f}%)")

        cap.release()
        out.release()

        print(f" Video processing complete! Created annotated video: {output_path}")
        print(f" Total detections: {len(detections)}")

        if not detections:
            return (
                f"🔍 No objects detected matching your criteria. Try lowering confidence threshold or selecting different objects.",
                pd.DataFrame(),
                """<div style="background: linear-gradient(135deg, #f59e0b, #d97706); 
                           color: white; padding: 20px; border-radius: 15px; text-align: center;">
                    ⚠️ No Detections Found - Annotated video shows original footage
                   </div>""",
                gr.update(visible=False),
                video_file,  # Original video
                output_path  # Annotated video (even if no detections)
            )

        # Create DataFrame
        df = pd.DataFrame(detections)

        # Results summary
        unique_objects = df['object'].nunique()
        total_detections = len(df)

        results_html = f"""
        <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); 
                    color: white; padding: 25px; border-radius: 20px; margin: 20px 0;">
            <h3 style="text-align: center; margin-bottom: 20px;">🎯 YOLOv8 Analysis Complete</h3>

            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin-bottom: 20px;">
                <div style="background: rgba(34, 197, 94, 0.2); padding: 15px; border-radius: 10px; text-align: center;">
                    <div style="font-size: 2em; font-weight: bold; color: #22c55e;">{total_detections}</div>
                    <div>Total Detections</div>
                </div>
                <div style="background: rgba(59, 130, 246, 0.2); padding: 15px; border-radius: 10px; text-align: center;">
                    <div style="font-size: 2em; font-weight: bold; color: #3b82f6;">{unique_objects}</div>
                    <div>Object Types</div>
                </div>
                <div style="background: rgba(168, 85, 247, 0.2); padding: 15px; border-radius: 10px; text-align: center;">
                    <div style="font-size: 2em; font-weight: bold; color: #a855f7;">{frame_count}</div>
                    <div>Frames Processed</div>
                </div>
            </div>

            <div style="background: linear-gradient(135deg, #22c55e, #16a34a); 
                        padding: 15px; border-radius: 10px; text-align: center; color: white;">
                <span style="font-size: 1.1em;">✅ Annotated video with bounding boxes is ready for viewing!</span>
            </div>
        </div>
        """

        return (
            f"✅ Analysis Complete! Found {total_detections} detections across {unique_objects} object types.",
            df,
            results_html,
            gr.update(visible=True),
            video_file,  # Original video
            output_path  # Annotated video
        )

    except Exception as e:
        error_msg = f" Error: {str(e)}"
        print(f"ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        return (
            error_msg,
            pd.DataFrame(),
            f"""<div style="background: #dc2626; color: white; padding: 20px; border-radius: 15px;">
                 Processing Error: {str(e)}
                </div>""",
            gr.update(visible=False),
            video_file,  # Original video
            None  # Annotated video
        )


def generate_intelligence_report(analysis_results_df, video_file_name="surveillance_footage.mp4", fps=30):
    """Generate comprehensive intelligence report with navy-themed styling"""
    if analysis_results_df.empty:
        return (
            " No analysis results available. Please perform object detection first.",
            "",
            gr.update(visible=False),
            gr.update(visible=False)
        )

    time.sleep(2)

    total_detections = len(analysis_results_df)
    unique_objects = analysis_results_df['object'].nunique()
    object_counts = analysis_results_df['object'].value_counts().to_dict()
    threat_distribution = analysis_results_df['threat_level'].value_counts().to_dict()

    confidence_values = [float(conf.rstrip('%')) for conf in analysis_results_df['confidence']]
    avg_confidence = sum(confidence_values) / len(confidence_values) if confidence_values else 0
    max_confidence = max(confidence_values) if confidence_values else 0
    min_confidence = min(confidence_values) if confidence_values else 0

    highlight_color = "#101d6b"
    background_color = "#0f172a"
    card_text_color = "white"

    intelligence_html = f"""
    <div style="background: {background_color}; color: white; padding: 30px; border-radius: 25px; margin: 20px 0;
                box-shadow: 0 20px 50px rgba(0,0,0,0.4);">

        <div style="text-align: center; margin-bottom: 30px;">
            <h2 style="margin: 0; font-size: 2.2em;"> MILITARY SURVEILLANCE INTELLIGENCE REPORT</h2>
            <p style="margin: 10px 0 0 0; font-size: 1.1em;">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
            <div style="background: rgba(255,255,255,0.05); padding: 10px; border-radius: 10px; margin: 15px 0;">
                <strong>Source Material:</strong> {video_file_name} | <strong>AI Model:</strong> YOLOv8 Custom Trained
            </div>
        </div>

        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin-bottom: 30px;">
            <div style="background: {highlight_color}; color: {card_text_color}; padding: 25px; border-radius: 20px; text-align: center;">
                <h3 style="margin: 0 0 15px 0;"> TOTAL DETECTIONS</h3>
                <div style="font-size: 3em; margin: 10px 0; font-weight: bold;">{total_detections}</div>
                <p style="margin: 0;">Detection Events</p>
            </div>
            <div style="background: {highlight_color}; color: {card_text_color}; padding: 25px; border-radius: 20px; text-align: center;">
                <h3 style="margin: 0 0 15px 0;">📊 OBJECT TYPES</h3>
                <div style="font-size: 3em; margin: 10px 0; font-weight: bold;">{unique_objects}</div>
                <p style="margin: 0;">Categories Found</p>
            </div>
            <div style="background: {highlight_color}; color: {card_text_color}; padding: 25px; border-radius: 20px; text-align: center;">
                <h3 style="margin: 0 0 15px 0;">📈 AVG CONFIDENCE</h3>
                <div style="font-size: 3em; margin: 10px 0; font-weight: bold;">{avg_confidence:.1f}%</div>
                <p style="margin: 0;">Detection Accuracy</p>
            </div>
        </div>

        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 25px; margin-bottom: 25px;">
            <div style="background: rgba(255,255,255,0.05); padding: 25px; border-radius: 20px;">
                <h3 style="margin: 0 0 20px 0;">⚠️ THREAT ASSESSMENT</h3>
    """

    for threat, count in threat_distribution.items():
        intelligence_html += f"""
                <div style="display: flex; justify-content: space-between; align-items: center; 
                            padding: 12px; background: {highlight_color}; color: white; border-radius: 10px; margin-bottom: 10px;">
                    <span style="font-weight: bold;">{threat.upper()}</span>
                    <span style="font-size: 1.3em; font-weight: bold;">{count}</span>
                </div>
        """

    intelligence_html += f"""
            </div>

            <div style="background: rgba(255,255,255,0.05); padding: 25px; border-radius: 20px;">
                <h3 style="margin: 0 0 20px 0;">📦 OBJECT BREAKDOWN</h3>
                <div style="max-height: 200px; overflow-y: auto;">
    """

    for obj, count in object_counts.items():
        intelligence_html += f"""
                    <div style="display: flex; justify-content: space-between; padding: 8px 0; 
                                border-bottom: 1px solid rgba(255,255,255,0.2);">
                        <span>{obj}</span>
                        <span style="font-weight: bold; color: {highlight_color};">{count}</span>
                    </div>
        """

    intelligence_html += f"""
                </div>
            </div>
        </div>

        <div style="background: rgba(255,255,255,0.03); padding: 20px; border-radius: 15px;">
            <h3 style="margin: 0 0 15px 0;">📊 STATISTICAL SUMMARY</h3>
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
                <div><strong>Highest Confidence:</strong> {max_confidence:.1f}%</div>
                <div><strong>Lowest Confidence:</strong> {min_confidence:.1f}%</div>
                <div><strong>Detection Range:</strong> {max_confidence - min_confidence:.1f}%</div>
                <div><strong>Processing Status:</strong> Complete</div>
            </div>
        </div>
    </div>
    """

    return (
        "✅ Intelligence report generated successfully! Download reports below.",
        intelligence_html,
        gr.update(visible=True),
        gr.update(visible=True)
    )


def create_json_report(analysis_results_df, video_filename="surveillance_footage.mp4"):
    """Generate downloadable JSON intelligence report"""
    if analysis_results_df.empty:
        return None

    # Calculate statistics
    total_detections = len(analysis_results_df)
    unique_objects = analysis_results_df['object'].nunique()
    object_counts = analysis_results_df['object'].value_counts().to_dict()
    threat_distribution = analysis_results_df['threat_level'].value_counts().to_dict()

    # Get confidence statistics
    confidence_values = []
    for conf in analysis_results_df['confidence']:
        if isinstance(conf, str):
            conf_num = float(conf.rstrip('%'))
        else:
            conf_num = float(conf) * 100
        confidence_values.append(conf_num)

    avg_confidence = sum(confidence_values) / len(confidence_values) if confidence_values else 0
    max_confidence = max(confidence_values) if confidence_values else 0
    min_confidence = min(confidence_values) if confidence_values else 0

    # Create JSON report structure
    report_data = {
        "report_metadata": {
            "report_id": f"SURV-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            "generated_timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "source_video": video_filename,
            "analysis_system": "YOLOv8 Custom Military Object Detection",
            "report_version": "v2.1",
            "classification": "CONFIDENTIAL"
        },
        "executive_summary": {
            "total_detections": total_detections,
            "unique_object_types": unique_objects,
            "average_confidence": round(avg_confidence, 2),
            "highest_confidence": round(max_confidence, 1),
            "lowest_confidence": round(min_confidence, 1),
            "confidence_range": round(max_confidence - min_confidence, 1),
            "primary_threat_level": max(threat_distribution,
                                        key=threat_distribution.get) if threat_distribution else "None"
        },
        "threat_assessment": {
            "threat_distribution": threat_distribution,
            "critical_threats": threat_distribution.get('Critical', 0),
            "high_threats": threat_distribution.get('High', 0),
            "medium_threats": threat_distribution.get('Medium', 0),
            "low_threats": threat_distribution.get('Low', 0),
            "total_high_priority": threat_distribution.get('Critical', 0) + threat_distribution.get('High', 0)
        },
        "object_analysis": {
            "object_counts": object_counts,
            "most_detected_object": max(object_counts, key=object_counts.get) if object_counts else "None",
            "most_detected_count": max(object_counts.values()) if object_counts else 0,
            "object_diversity_index": round((unique_objects / total_detections) * 100, 1) if total_detections > 0 else 0
        },
        "detailed_detections": [],
        "statistical_analysis": {
            "confidence_metrics": {
                "average": round(avg_confidence, 2),
                "maximum": round(max_confidence, 1),
                "minimum": round(min_confidence, 1),
                "range": round(max_confidence - min_confidence, 1),
                "reliability_rating": "HIGH" if avg_confidence > 70 else "MEDIUM" if avg_confidence > 50 else "LOW"
            },
            "detection_patterns": {
                "total_events": total_detections,
                "objects_per_category": object_counts,
                "threat_level_distribution": threat_distribution
            }
        },
        "recommendations": {
            "immediate_actions": [
                "Deploy additional surveillance to high-threat coordinates",
                "Establish continuous monitoring protocols",
                "Coordinate with ground intelligence units"
            ],
            "strategic_recommendations": [
                f"Pattern analysis suggests {'organized military presence' if total_detections > 50 else 'limited military activity'}",
                f"Recommend {'enhanced' if threat_distribution.get('Critical', 0) > 0 else 'standard'} security protocols",
                f"Consider {'immediate' if threat_distribution.get('Critical', 0) > 0 else 'routine'} response measures"
            ],
            "priority_level": "HIGH" if threat_distribution.get('Critical',
                                                                0) > 0 else "MEDIUM" if threat_distribution.get('High',
                                                                                                                0) > 0 else "LOW"
        }
    }

    # Add detailed detection records
    for i, (_, row) in enumerate(analysis_results_df.iterrows(), 1):
        confidence_value = row['confidence']
        if isinstance(confidence_value, str):
            confidence_clean = float(confidence_value.rstrip('%'))
        else:
            confidence_clean = float(confidence_value) * 100

        detection_record = {
            "detection_id": f"DET-{i:03d}",
            "object_type": row['object'],
            "timestamp": row['timestamp'],
            "duration": row['duration'],
            "confidence_percentage": round(confidence_clean, 1),
            "threat_level": row['threat_level'],
            "coordinates": row['coordinates'],
            "detection_count": row['count']
        }
        report_data["detailed_detections"].append(detection_record)

    # Create temporary JSON file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json',
                                     prefix=f'Intelligence_Report_{timestamp}_') as f:
        json.dump(report_data, f, indent=2, ensure_ascii=False)
        temp_path = f.name

    return temp_path


def create_pdf_report(analysis_df, video_filename="surveillance_footage.mp4"):
    """Generate user-friendly PDF report without Unicode characters"""
    if analysis_df.empty or not PDF_AVAILABLE:
        return None

    try:
        class PDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 16)
                # Remove emojis from header
                self.cell(0, 10, 'MILITARY SURVEILLANCE INTELLIGENCE REPORT', 0, 1, 'C')
                self.ln(5)

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        pdf = PDF()
        pdf.add_page()

        # Header information
        pdf.set_font("Helvetica", '', 11)
        pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1)
        pdf.cell(0, 8, f"Source Video: {video_filename}", 0, 1)
        pdf.cell(0, 8, f"AI System: YOLOv8 Custom Military Detection Model", 0, 1)
        pdf.ln(10)

        # Summary section
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "EXECUTIVE SUMMARY", 0, 1)
        pdf.set_font("Helvetica", '', 11)

        total_detections = len(analysis_df)
        unique_objects = analysis_df['object'].nunique()

        # Get confidence values
        confidence_values = []
        for conf in analysis_df['confidence']:
            if isinstance(conf, str):
                conf_num = float(conf.rstrip('%'))
            else:
                conf_num = float(conf) * 100
            confidence_values.append(conf_num)

        avg_confidence = sum(confidence_values) / len(confidence_values) if confidence_values else 0

        pdf.cell(0, 8, f"- Total Detection Events: {total_detections}", 0, 1)
        pdf.cell(0, 8, f"- Unique Object Types Found: {unique_objects}", 0, 1)
        pdf.cell(0, 8, f"- Average Detection Confidence: {avg_confidence:.1f}%", 0, 1)
        pdf.ln(5)

        # Threat Assessment
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "THREAT ASSESSMENT", 0, 1)
        pdf.set_font("Helvetica", '', 11)

        threat_counts = analysis_df['threat_level'].value_counts().to_dict()

        # Define threat indicators without emojis
        threat_indicators = {
            "Critical": "[CRITICAL]",
            "High": "[HIGH]",
            "Medium": "[MEDIUM]",
            "Low": "[LOW]"
        }

        for threat_level, count in threat_counts.items():
            indicator = threat_indicators.get(threat_level, "[UNKNOWN]")
            pdf.cell(0, 8, f"{indicator} {threat_level} Level Threats: {count}", 0, 1)
        pdf.ln(5)

        # Object Breakdown
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "DETECTED OBJECTS", 0, 1)
        pdf.set_font("Helvetica", '', 11)

        object_counts = analysis_df['object'].value_counts().to_dict()
        for obj_type, count in object_counts.items():
            # Handle potential encoding issues
            obj_text = str(obj_type).encode('latin-1', 'ignore').decode('latin-1')
            pdf.cell(0, 8, f"- {obj_text}: {count} detections", 0, 1)
        pdf.ln(10)

        # Detailed Detection Table
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "DETAILED DETECTION LOG", 0, 1)
        pdf.set_font("Helvetica", 'B', 9)

        # Table headers
        pdf.cell(35, 8, "Object Type", 1, 0, 'C')
        pdf.cell(25, 8, "Time", 1, 0, 'C')
        pdf.cell(20, 8, "Confidence", 1, 0, 'C')
        pdf.cell(25, 8, "Threat Level", 1, 0, 'C')
        pdf.cell(20, 8, "Duration", 1, 0, 'C')
        pdf.cell(65, 8, "Coordinates", 1, 1, 'C')

        pdf.set_font("Helvetica", '', 8)

        # Limit rows to prevent page overflow
        max_rows = 25  # Limit to first 25 detections
        rows_added = 0

        for _, row in analysis_df.iterrows():
            if rows_added >= max_rows:
                pdf.ln(5)
                pdf.set_font("Helvetica", 'I', 9)
                pdf.cell(0, 8, f"... and {len(analysis_df) - max_rows} more detections", 0, 1)
                break

            # Truncate and clean text to fit in cells
            obj_text = str(row['object']).encode('latin-1', 'ignore').decode('latin-1')[:20]
            time_text = str(row['timestamp'])
            conf_text = str(row['confidence'])
            threat_text = str(row['threat_level'])[:15]
            dur_text = str(row['duration'])
            coord_text = str(row['coordinates'])[:30]

            pdf.cell(35, 6, obj_text, 1, 0)
            pdf.cell(25, 6, time_text, 1, 0)
            pdf.cell(20, 6, conf_text, 1, 0)
            pdf.cell(25, 6, threat_text, 1, 0)
            pdf.cell(20, 6, dur_text, 1, 0)
            pdf.cell(65, 6, coord_text, 1, 1)

            rows_added += 1

        # Recommendations
        pdf.ln(10)
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "RECOMMENDATIONS", 0, 1)
        pdf.set_font("Helvetica", '', 11)

        critical_count = threat_counts.get('Critical', 0)
        high_count = threat_counts.get('High', 0)

        if critical_count > 0:
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 8, "[IMMEDIATE ACTION REQUIRED]", 0, 1)
            pdf.set_font("Helvetica", '', 11)
            pdf.cell(0, 8, "- Deploy emergency response teams immediately", 0, 1)
            pdf.cell(0, 8, "- Establish secure perimeter around threat locations", 0, 1)
            pdf.cell(0, 8, "- Contact military command for further instructions", 0, 1)
            pdf.cell(0, 8, "- Evacuate civilian personnel from high-risk areas", 0, 1)
        elif high_count > 0:
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 8, "[HIGH PRIORITY MONITORING]", 0, 1)
            pdf.set_font("Helvetica", '', 11)
            pdf.cell(0, 8, "- Increase surveillance frequency in detected areas", 0, 1)
            pdf.cell(0, 8, "- Prepare rapid response protocols", 0, 1)
            pdf.cell(0, 8, "- Coordinate with intelligence units", 0, 1)
        else:
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 8, "[STANDARD MONITORING]", 0, 1)
            pdf.set_font("Helvetica", '', 11)
            pdf.cell(0, 8, "- Continue routine surveillance operations", 0, 1)
            pdf.cell(0, 8, "- Maintain current security alert levels", 0, 1)
            pdf.cell(0, 8, "- Regular reporting to command structure", 0, 1)

        # Add new page for conclusion if needed
        if pdf.get_y() > 250:  # Near bottom of page
            pdf.add_page()

        # Conclusion
        pdf.ln(10)
        pdf.set_font("Helvetica", 'B', 14)
        pdf.cell(0, 10, "CONCLUSION", 0, 1)
        pdf.set_font("Helvetica", '', 11)

        conclusion_lines = [
            f"This surveillance analysis identified {total_detections} detection events",
            f"across {unique_objects} different object categories with an average",
            f"confidence level of {avg_confidence:.1f}%.",
            "",
            f"The threat assessment indicates {threat_counts.get('Critical', 0)} critical threats,",
            f"{threat_counts.get('High', 0)} high-priority threats, and {threat_counts.get('Medium', 0)} medium-level",
            "concerns that require appropriate response protocols.",
            "",
            "This report should be shared with relevant military personnel",
            "and used for strategic decision-making regarding the surveillance area.",
            "",
            "CLASSIFICATION: CONFIDENTIAL",
            "DISTRIBUTION: AUTHORIZED PERSONNEL ONLY"
        ]

        for line in conclusion_lines:
            pdf.cell(0, 6, line, 0, 1)

        # Save PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(tempfile.gettempdir(), f"Military_Surveillance_Report_{timestamp}.pdf")
        pdf.output(file_path)
        return file_path

    except Exception as e:
        print(f"❌ Error creating PDF: {e}")
        import traceback
        traceback.print_exc()
        return None


def create_docx_report(analysis_df, video_filename="surveillance_footage.mp4"):
    """Generate user-friendly Word document report"""
    if analysis_df.empty or not DOCX_AVAILABLE:
        return None

    try:
        doc = Document()

        # Title
        title = doc.add_heading('Military Surveillance Intelligence Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        subtitle_run.bold = True

        # Video info
        doc.add_paragraph(f"Source Video: {video_filename}")
        doc.add_paragraph(f"AI System: YOLOv8 Custom Military Object Detection Model")
        doc.add_paragraph(f"Model Classes: Artillery, Missile, Radar, M. Rocket Launcher, Soldier, Tank, Vehicle")
        doc.add_paragraph("")

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)

        total_detections = len(analysis_df)
        unique_objects = analysis_df['object'].nunique()

        # Calculate confidence
        confidence_values = []
        for conf in analysis_df['confidence']:
            if isinstance(conf, str):
                conf_num = float(conf.rstrip('%'))
            else:
                conf_num = float(conf) * 100
            confidence_values.append(conf_num)

        avg_confidence = sum(confidence_values) / len(confidence_values) if confidence_values else 0

        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Detection Events: ").bold = True
        summary_para.add_run(f"{total_detections}\n")
        summary_para.add_run(f"Unique Object Types: ").bold = True
        summary_para.add_run(f"{unique_objects}\n")
        summary_para.add_run(f"Average Confidence: ").bold = True
        summary_para.add_run(f"{avg_confidence:.1f}%\n")

        # Add explanation for general audience
        explanation = doc.add_paragraph()
        explanation.add_run("What this means: ").bold = True
        explanation.add_run(
            "Our AI system analyzed the video and identified military objects like tanks, soldiers, and weapons. ")
        explanation.add_run(
            f"It found {total_detections} different objects and is {avg_confidence:.0f}% confident about the detections.")

        # Threat Assessment
        doc.add_heading('Threat Assessment', level=1)
        threat_counts = analysis_df['threat_level'].value_counts().to_dict()

        # Simple explanation
        threat_para = doc.add_paragraph()
        threat_para.add_run("Threat Levels Explained:\n").bold = True
        threat_para.add_run("• Critical = Very Dangerous - Need immediate response\n")
        threat_para.add_run("• High = Dangerous - Need careful monitoring\n")
        threat_para.add_run("• Medium = Somewhat dangerous - Need regular surveillance\n")
        threat_para.add_run("• Low = Not very dangerous - Normal monitoring\n\n")

        threat_table = doc.add_table(rows=1, cols=2)
        threat_table.style = 'Light Grid Accent 1'
        hdr_cells = threat_table.rows[0].cells
        hdr_cells[0].text = 'Threat Level'
        hdr_cells[1].text = 'Count'

        for threat_level, count in threat_counts.items():
            row_cells = threat_table.add_row().cells
            threat_indicator = "[CRITICAL]" if threat_level == "Critical" else "[HIGH]" if threat_level == "High" else "[MEDIUM]" if threat_level == "Medium" else "[LOW]"
            row_cells[0].text = f"{threat_indicator} {threat_level}"
            row_cells[1].text = str(count)

        doc.add_paragraph("")

        # Object Detection Results
        doc.add_heading('Detected Objects', level=1)
        object_counts = analysis_df['object'].value_counts().to_dict()

        # Map model class names to friendly descriptions
        friendly_names = {
            'Tank': 'Army Tank (Heavy armored vehicle)',
            'Soldier': 'Soldier (Military personnel)',
            'Vehicle': 'Military Vehicle (Army truck/car)',
            'Artillery': 'Artillery (Large gun/cannon)',
            'Missile': 'Missile (Guided weapon)',
            'Radar': 'Radar (Detection equipment)',
            'M. Rocket Launcher': 'Rocket Launcher (Weapon system)'
        }

        for obj_type, count in object_counts.items():
            friendly_name = friendly_names.get(obj_type, f"{obj_type}")
            doc.add_paragraph(f"• {friendly_name}: {count} detections", style='List Bullet')

        # Detailed Detection Table (limit to first 20 for readability)
        doc.add_heading('Detailed Detection Log (First 20 detections)', level=1)

        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Object Type'
        hdr_cells[1].text = 'Time'
        hdr_cells[2].text = 'Confidence'
        hdr_cells[3].text = 'Threat Level'
        hdr_cells[4].text = 'Duration'
        hdr_cells[5].text = 'Coordinates'

        # Add data rows (limit to first 20)
        row_count = 0
        for _, row in analysis_df.iterrows():
            if row_count >= 20:
                break
            row_cells = table.add_row().cells

            # Use friendly names
            friendly_obj = friendly_names.get(str(row['object']), str(row['object']))
            row_cells[0].text = friendly_obj
            row_cells[1].text = str(row['timestamp'])
            row_cells[2].text = str(row['confidence'])
            row_cells[3].text = str(row['threat_level'])
            row_cells[4].text = str(row['duration'])
            row_cells[5].text = str(row['coordinates'])
            row_count += 1

        if len(analysis_df) > 20:
            doc.add_paragraph(f"... and {len(analysis_df) - 20} more detections not shown here.")

        # Recommendations
        doc.add_heading('Recommended Actions', level=1)

        critical_count = threat_counts.get('Critical', 0)
        high_count = threat_counts.get('High', 0)

        if critical_count > 0:
            rec_para = doc.add_paragraph()
            rec_para.add_run("IMMEDIATE ACTION REQUIRED:\n").bold = True
            doc.add_paragraph("• Deploy emergency response teams immediately", style='List Bullet')
            doc.add_paragraph("• Establish secure perimeter around threat locations", style='List Bullet')
            doc.add_paragraph("• Contact military command for further instructions", style='List Bullet')
            doc.add_paragraph("• Evacuate civilian personnel from high-risk areas", style='List Bullet')
        elif high_count > 0:
            rec_para = doc.add_paragraph()
            rec_para.add_run("HIGH PRIORITY MONITORING:\n").bold = True
            doc.add_paragraph("• Increase surveillance frequency in detected areas", style='List Bullet')
            doc.add_paragraph("• Prepare rapid response protocols", style='List Bullet')
            doc.add_paragraph("• Coordinate with intelligence teams", style='List Bullet')
        else:
            rec_para = doc.add_paragraph()
            rec_para.add_run("STANDARD MONITORING:\n").bold = True
            doc.add_paragraph("• Continue routine surveillance operations", style='List Bullet')
            doc.add_paragraph("• Maintain current security alert levels", style='List Bullet')
            doc.add_paragraph("• Send regular reports to command structure", style='List Bullet')

        # Simple Conclusion
        doc.add_heading('Summary', level=1)
        conclusion_text = f"""
Analysis Summary: Our AI system analyzed surveillance footage and identified {total_detections} military objects.

Key Findings: We detected {unique_objects} different types of military equipment with an average confidence level of {avg_confidence:.0f}%.

Threat Assessment: Found {threat_counts.get('Critical', 0)} critical threats, {threat_counts.get('High', 0)} high-priority threats, and {threat_counts.get('Medium', 0)} medium-level concerns.

Recommendation: This information should be used by military personnel to make informed security decisions for the surveilled area.

Report Classification: This report should be shared with authorized military personnel and security decision-makers only.
        """
        doc.add_paragraph(conclusion_text.strip())

        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(tempfile.gettempdir(), f"Military_Surveillance_Report_{timestamp}.docx")
        doc.save(file_path)
        return file_path

    except Exception as e:
        print(f"❌ Error creating Word document: {e}")
        return None


# Enhanced CSS
military_css = """
.gradio-container {
    background: linear-gradient(135deg, #0f172a 0%, #1e293b 50%, #334155 100%);
    font-family: 'Segoe UI', 'Consolas', monospace;
    color: #e2e8f0;
}

.main-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    color: #ffffff;
    padding: 35px !important;
    border-radius: 20px;
    text-align: center;
    margin-bottom: 25px;
    box-shadow: 0 15px 35px rgba(0,0,0,0.4);
    border: 2px solid #334155;
}

.main-header h1 {
    margin: 0;
    font-size: 2.9em !important;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.7);
}

.section-header {
    background: linear-gradient(135deg, #841616 0%, #b91c1c 100%);
    color: white;
    padding: 18px 22px;
    border-radius: 15px;
    margin: 15px 0;
    font-weight: bold;
    text-transform: uppercase;
    letter-spacing: 1px;
    border: 1px solid #ef4444;
    font-size: 1.1em;
}
"""


def create_yolov8_surveillance_interface():
    """Create the main Gradio interface for YOLOv8 surveillance analysis"""

    with gr.Blocks(
            css=military_css,
            title="Military Surveillance Analyzer",
            theme=gr.themes.Base(primary_hue="red", secondary_hue="slate", neutral_hue="slate")
    ) as demo:
        # Header
        gr.HTML(f"""
        <div class="main-header">
            <h1>MILITARY SURVEILLANCE ANALYZER</h1>
            <p style="margin: 10px 0 0 0; font-size: 1.1em; opacity: 0.9; color: #fbbf24;">
                {'CUSTOM MODEL LOADED' if MODEL_LOADED else '🟡 DEMO MODE - MOCK DETECTIONS'} | AI-Powered Object Detection
            </p>
        </div>
        """)

        # State variable to store analysis results
        analysis_data_state = gr.State(pd.DataFrame())

        # Main interface
        with gr.Tab(" YOLOv8 Object Detection"):
            with gr.Row():
                with gr.Column(scale=1):
                    gr.HTML('<div class="section-header"> Upload Surveillance Footage</div>')

                    video_input = gr.File(
                        label="Upload Video File",
                        file_types=["video"],
                        file_count="single"
                    )

                    gr.HTML('<div class="section-header"> Select Detection Targets</div>')

                    object_selector = gr.CheckboxGroup(
                        choices=DETECTION_CATEGORIES,
                        label="Objects to Detect",
                        value=DETECTION_CATEGORIES,  # Pre-select ALL items to catch detections
                        interactive=True
                    )

                    confidence_threshold = gr.Slider(
                        label="Confidence Threshold (%)",
                        minimum=1,  # Lower minimum to catch more detections
                        maximum=95,
                        value=15,  # Lower default value
                        step=1,
                        interactive=True
                    )

                    analyze_btn = gr.Button(
                        " Analyze with YOLOv8",
                        variant="primary",
                        size="lg"
                    )

                with gr.Column(scale=2):
                    analysis_status = gr.Textbox(
                        label="Analysis Status",
                        interactive=False,
                        placeholder="Upload video and select objects to analyze..."
                    )

                    # Create two video displays side by side
                    with gr.Row():
                        with gr.Column():
                            original_video_display = gr.Video(
                                label=" Original Video",
                                height=300,
                                show_label=True
                            )

                        with gr.Column():
                            annotated_video_display = gr.Video(
                                label=" Annotated Video (YOLOv8 Detections)",
                                height=300,
                                show_label=True
                            )

                    results_display = gr.HTML()

                    detection_results = gr.Dataframe(
                        headers=["Object", "Timestamp", "Duration", "Confidence", "Count", "Threat Level",
                                 "Coordinates"],
                        label=" YOLOv8 Detection Results",
                        interactive=False
                    )

        # Intelligence reporting
        with gr.Tab(" Intelligence Report"):
            with gr.Column():
                generate_report_btn = gr.Button(
                    " Generate Intelligence Report",
                    variant="secondary",
                    size="lg",
                    visible=False
                )

                report_status = gr.Textbox(
                    label="Report Status",
                    interactive=False,
                    placeholder="Complete analysis first..."
                )

                intelligence_display = gr.HTML()

                # Download buttons for different formats
                gr.HTML('<div style="margin: 20px 0;"><h3> Download Reports</h3></div>')

                with gr.Row():
                    download_btn_json = gr.DownloadButton(
                        " Download JSON Report",
                        variant="primary",
                        size="lg",
                        visible=False
                    )
                    download_btn_pdf = gr.DownloadButton(
                        " Download PDF Report",
                        variant="primary",
                        size="lg",
                        visible=False
                    )
                    download_btn_docx = gr.DownloadButton(
                        " Download Word Report",
                        variant="primary",
                        size="lg",
                        visible=False
                    )

        # System info
        with gr.Tab(" System Information"):
            gr.HTML(f"""
            <div style="background: #202020; color: white; padding: 30px; border-radius: 20px;">

                <h2>YOLOv8 Military Surveillance System</h2>

                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; margin: 20px 0;">
                    <h3>Model Status</h3>
                    <p><strong>Model Path:</strong> {MODEL_PATH}</p>
                    <p><strong>Status:</strong> {'✅ Loaded Successfully' if MODEL_LOADED else '❌ Not Available (Using Mock Data)'}</p>
                    <p><strong>Framework:</strong> YOLOv8 (Ultralytics)</p>
                    <p><strong>Export Formats:</strong> JSON, PDF{'✅' if PDF_AVAILABLE else '❌'}, Word{'✅' if DOCX_AVAILABLE else '❌'}</p>
                </div>

                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; margin: 20px 0;">
                    <h3>Capabilities</h3>
                    <ul>
                        <li>Real-time object detection in surveillance footage</li>
                        <li>Custom trained model for military objects</li>
                        <li>Threat level assessment and classification</li>
                        <li>Temporal analysis with precise timestamps</li>
                        <li>Batch processing with confidence filtering</li>
                        <li>Multi-format intelligence reports (JSON, PDF, Word)</li>
                        <li>Side-by-side original and annotated video display</li>
                    </ul>
                </div>

                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px;">
                    <h3>Usage Instructions</h3>
                    <ol>
                        <li>Upload your surveillance video file</li>
                        <li>Select object types you want to detect</li>
                        <li>Set confidence threshold (15% recommended for maximum detection)</li>
                        <li>Click "Analyze with YOLOv8" to process</li>
                        <li>View both original and annotated videos side by side</li>
                        <li>Generate and download intelligence reports in multiple formats</li>
                    </ol>
                </div>

                <div style="background: rgba(255,255,255,0.1); padding: 20px; border-radius: 15px; margin: 20px 0;">
                    <h3>Report Formats</h3>
                    <ul>
                        <li><strong>JSON:</strong> Technical data for systems integration</li>
                        <li><strong>PDF:</strong> Professional printable report (emoji-free)</li>  
                        <li><strong>Word:</strong> Editable document with tables and formatting</li>
                    </ul>
                </div>
            </div>
            """)

        # Event handlers
        # Update original video display when file is uploaded
        video_input.change(
            fn=lambda x: x,
            inputs=[video_input],
            outputs=[original_video_display]
        )

        def analyze_and_store(video_file, selected_objects, confidence_threshold):
            """Analyze video and store results in state"""
            result = analyze_surveillance_footage(video_file, selected_objects, confidence_threshold)
            analysis_status, detection_df, results_html, report_btn_update, original_video, annotated_video = result

            # Store the DataFrame in state
            return (
                analysis_status,
                detection_df,
                results_html,
                report_btn_update,
                original_video,  # Original video
                annotated_video,  # Annotated video
                detection_df  # This goes to the state
            )

        analyze_btn.click(
            fn=analyze_and_store,
            inputs=[video_input, object_selector, confidence_threshold],
            outputs=[analysis_status, detection_results, results_display, generate_report_btn,
                     original_video_display, annotated_video_display, analysis_data_state]
        )

        def generate_report_and_download(analysis_df, video_input_file):
            """Generate report and prepare downloads"""
            if analysis_df.empty:
                return (
                    "❌ No analysis results available. Please perform object detection first.",
                    "",
                    gr.update(visible=False),
                    None,
                    None,
                    None
                )

            video_name = video_input_file.name if video_input_file else "surveillance_footage.mp4"

            # Generate HTML report
            report_result = generate_intelligence_report(analysis_df, video_name)
            status, html, _, _ = report_result

            # Generate files for download
            json_file_path = create_json_report(analysis_df, video_name)
            pdf_file_path = create_pdf_report(analysis_df, video_name)
            docx_file_path = create_docx_report(analysis_df, video_name)

            return (
                status,
                html,
                gr.update(visible=True),  # JSON button
                json_file_path,
                pdf_file_path,
                docx_file_path
            )

        generate_report_btn.click(
            fn=generate_report_and_download,
            inputs=[analysis_data_state, video_input],
            outputs=[
                report_status,
                intelligence_display,
                download_btn_json,
                download_btn_json,
                download_btn_pdf,
                download_btn_docx
            ]
        )

        # Update visibility of PDF and Word buttons after generation
        def show_download_buttons():
            return (
                gr.update(visible=True),  # JSON
                gr.update(visible=True),  # PDF
                gr.update(visible=True)  # Word
            )

        generate_report_btn.click(
            fn=show_download_buttons,
            outputs=[download_btn_json, download_btn_pdf, download_btn_docx]
        )

    return demo


# Launch the application
if __name__ == "__main__":
    demo = create_yolov8_surveillance_interface()

    print("🛡️ Starting YOLOv8 Military Surveillance Analyzer...")
    print("=" * 60)
    print(f"📍 Model Status: {'✅ Custom YOLOv8 Loaded' if MODEL_LOADED else '⚠️ Demo Mode'}")
    print(f"📁 Model Path: {MODEL_PATH}")
    print(f"📄 PDF Support: {'✅ Available' if PDF_AVAILABLE else '❌ Install fpdf2'}")
    print(f"📝 Word Support: {'✅ Available' if DOCX_AVAILABLE else '❌ Install python-docx'}")

    demo.launch(
        server_name="127.0.0.1",
        server_port=7860,
        show_error=True,
        share=False,  # Set to False to avoid connection issues
        inbrowser=True
    )
